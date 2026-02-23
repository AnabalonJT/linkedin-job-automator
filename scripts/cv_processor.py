"""
Procesador de Currículum Vitae

Extrae información de PDF/DOCX y la estructura en JSON
para uso como contexto en consultas IA
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, asdict


@dataclass
class CVExperience:
    """Experiencia laboral"""
    company: str
    position: str
    duration: str
    duration_years: float
    technologies: List[str]
    achievements: List[str]
    description: str


@dataclass
class CVEducation:
    """Educación"""
    institution: str
    degree: str
    field: str
    year_completed: Optional[str]


@dataclass
class CVContext:
    """Contexto estructurado del CV"""
    name: str
    title: str
    summary: str
    email: Optional[str]
    phone: Optional[str]
    
    experience: List[Dict[str, Any]]
    education: List[Dict[str, Any]]
    
    technical_skills: List[str]
    soft_skills: List[str]
    languages: List[str]
    certifications: List[str]
    
    total_years_experience: float
    raw_text: str  # Texto completo para contexto adicional


class CVProcessor:
    """Procesador principal de CV - Soporta múltiples CVs"""
    
    def __init__(self, cv_software_path: Optional[str] = None, cv_engineer_path: Optional[str] = None):
        """
        Inicializar procesador con dos CVs
        
        Args:
            cv_software_path: Ruta al CV de software
            cv_engineer_path: Ruta al CV de consultoría/engineering
        """
        self.cv_software_path = cv_software_path or os.getenv('CV_SOFTWARE_PATH', 'config/CV_software.pdf')
        self.cv_engineer_path = cv_engineer_path or os.getenv('CV_ENGINEER_PATH', 'config/CV_engineer.pdf')
        
        # Usar siempre curriculum_context.json unificado
        self.cache_path = os.getenv('CV_CONTEXT_CACHE', 'data/curriculum_context.json')
        
        self.context = None  # Contexto único cargado en memoria
        self.debug = os.getenv('DEBUG', 'False').lower() == 'true'
    
    def load_or_create(self) -> CVContext:
        """
        Carga contexto unificado (usa siempre curriculum_context.json)
        
        Returns:
            CVContext único y consolidado
        """
        # Intenta cargar desde cache
        if os.path.exists(self.cache_path):
            if self.debug:
                print(f"📋 Cargando CV desde cache: {self.cache_path}")
            self.context = self.load_from_cache(self.cache_path)
            return self.context
        
        # Si no existe cache, crear desde CV actual
        if self.debug:
            print(f"📋 Procesando CV de software: {self.cv_software_path}")
        
        self.context = self.process_cv_file(self.cv_software_path)
        self.save_to_cache(self.context, self.cache_path)
        return self.context
    
    def get_context(self) -> CVContext:
        """
        Obtiene contexto actual (cargado o carga si no existe)
        
        Returns:
            CVContext consolidado
        """
        if not self.context:
            self.load_or_create()
        return self.context
    
    def get_context_string(self) -> str:
        """
        Obtiene contexto como string para prompts IA
        
        Returns:
            String formateado para usar en prompts
        """
        context = self.get_context()
        return self.get_context_as_string(context)
    
    def process_cv_file(self, file_path: str) -> CVContext:
        """
        Procesar archivo CV (PDF o DOCX)
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            CVContext estructurado
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"CV no encontrado: {file_path}")
        
        # Detectar tipo de archivo
        if file_path.endswith('.pdf'):
            text = self._extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Formato no soportado: {file_path}. Use PDF o DOCX")
        
        if self.debug:
            print(f"✅ CV extraído ({len(text)} caracteres)")
        
        # Estructurar el contexto
        context = self._structure_context(text)
        self.context = context
        return context
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrae texto de PDF"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber no instalado. Ejecuta: pip install pdfplumber")
        
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            if self.debug:
                print(f"⚠️  Error extrayendo PDF: {e}")
            raise
        
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extrae texto de DOCX"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx no instalado. Ejecuta: pip install python-docx")
        
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            if self.debug:
                print(f"⚠️  Error extrayendo DOCX: {e}")
            raise
        
        return text
    
    def _structure_context(self, raw_text: str) -> CVContext:
        """
        Estructurar el CV extraído en secciones
        
        Args:
            raw_text: Texto completo del CV
        
        Returns:
            CVContext estructurado
        """
        
        # Extraer información básica
        name = self._extract_name(raw_text)
        email = self._extract_email(raw_text)
        phone = self._extract_phone(raw_text)
        title = self._extract_title(raw_text)
        summary = self._extract_summary(raw_text)
        
        # Extraer secciones
        experience = self._extract_experience(raw_text)
        education = self._extract_education(raw_text)
        technical_skills = self._extract_technical_skills(raw_text)
        soft_skills = self._extract_soft_skills(raw_text)
        languages = self._extract_languages(raw_text)
        certifications = self._extract_certifications(raw_text)
        
        # Calcular años de experiencia total
        total_years = self._calculate_total_years(experience)
        
        context = CVContext(
            name=name,
            title=title,
            summary=summary,
            email=email,
            phone=phone,
            experience=experience,
            education=education,
            technical_skills=technical_skills,
            soft_skills=soft_skills,
            languages=languages,
            certifications=certifications,
            total_years_experience=total_years,
            raw_text=raw_text
        )
        
        if self.debug:
            print(f"📋 Contexto estructurado:")
            print(f"   Nombre: {name}")
            print(f"   Título: {title}")
            print(f"   Email: {email}")
            print(f"   Años experiencia: {total_years}")
            print(f"   Skills técnicos: {len(technical_skills)}")
            print(f"   Experiencias: {len(experience)}")
            print(f"   Educación: {len(education)}")
        
        return context
    
    def _extract_name(self, text: str) -> str:
        """Extrae nombre del candidato"""
        # Buscar nombres comunes en primeras líneas
        lines = text.split('\n')[:10]
        text_start = '\n'.join(lines).lower()
        return lines[0].strip() if lines else "Candidato"
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extrae email"""
        match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extrae teléfono"""
        # Buscar patrones de teléfono
        match = re.search(r'[\+]?[0-9\s\-\(\)]{10,}', text)
        return match.group(0) if match else None
    
    def _extract_title(self, text: str) -> str:
        """Extrae título/posición actual"""
        # Buscar en primeras líneas
        lines = text.split('\n')
        for line in lines[:20]:
            line = line.strip()
            if line and len(line) > 5 and len(line) < 100:
                # Podría ser título
                if any(word in line.lower() for word in ['engineer', 'developer', 'manager', 'lead', 'senior', 'junior', 'architect', 'analyst']):
                    return line
        return "Profesional"
    
    def _extract_summary(self, text: str) -> str:
        """Extrae resumen/about"""
        # Buscar sección de resumen
        patterns = [
            r'(?:PROFILE|ABOUT|SUMMARY|RESUMEN|PERFIL)(.*?)(?:EXPERIENCE|EXPERIENCIA|EDUCATION|EDUCACIÓN)',
            r'(.*?)(?:EXPERIENCE|EXPERIENCIA)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                if len(summary) > 50:  # Solo si es texto significativo
                    return summary[:300]  # Limitar a 300 caracteres
        
        return ""
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extrae experiencia laboral"""
        experience = []
        
        # Buscar secciones de experiencia
        patterns = [
            r'(?:EXPERIENCE|EXPERIENCIA)(.*?)(?:EDUCATION|EDUCACIÓN|SKILLS|LANGUAGES|$)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                exp_text = match.group(1)
                
                # Dividir por líneas vacías o patrones de empresa
                entries = re.split(r'\n\s*\n', exp_text)
                
                for entry in entries:
                    if len(entry) > 20:  # Solo si tiene contenido
                        job = {
                            'company': self._extract_company_name(entry),
                            'position': self._extract_job_position(entry),
                            'duration': self._extract_duration(entry),
                            'duration_years': self._calculate_duration_years(self._extract_duration(entry)),
                            'technologies': self._extract_technologies_from_entry(entry),
                            'achievements': self._extract_achievements(entry),
                            'description': entry.strip()[:200]
                        }
                        if job['position'] or job['company']:
                            experience.append(job)
        
        return experience[:10]  # Últimos 10 trabajos
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extrae educación"""
        education = []
        
        # Buscar sección de educación
        pattern = r'(?:EDUCATION|EDUCACIÓN)(.*?)(?:SKILLS|COMPETENCIAS|LANGUAGES|CERTIFICATIONS|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            edu_text = match.group(1)
            entries = re.split(r'\n\s*\n', edu_text)
            
            for entry in entries:
                if len(entry) > 10:
                    edu = {
                        'institution': self._extract_institution(entry),
                        'degree': self._extract_degree(entry),
                        'field': self._extract_field(entry),
                        'year_completed': self._extract_year(entry)
                    }
                    if edu['degree'] or edu['institution']:
                        education.append(edu)
        
        return education
    
    def _extract_technical_skills(self, text: str) -> List[str]:
        """Extrae skills técnicos"""
        skills = []
        
        # Buscar sección de skills
        patterns = [
            r'(?:TECHNICAL SKILLS|TECHNICAL|SKILLS|COMPETENCIAS TÉCNICAS)(.*?)(?:SOFT SKILLS|LANGUAGES|CERTIFICATIONS|PROJECTS|$)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                skills_text = match.group(1)
                
                # Dividir por comas, puntos o líneas
                candidates = re.split(r'[,\n•\-]', skills_text)
                for skill in candidates:
                    skill = skill.strip()
                    if len(skill) > 2 and len(skill) < 50:
                        skills.append(skill)
        
        # Tecnologías comunes (como fallback)
        common_techs = ['Python', 'JavaScript', 'React', 'Angular', 'Django',
                       'FastAPI', 'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'PostgreSQL', 'MySQL', 'Redis', 'Git', 'Linux', 'Windows', 'macOS']
        
        for tech in common_techs:
            if tech.lower() in text.lower():
                skills.append(tech)
        
        return list(set(skills))[:30]  # Unique, máx 30
    
    def _extract_soft_skills(self, text: str) -> List[str]:
        """Extrae soft skills"""
        soft_skills = []
        
        # Palabras clave de soft skills
        soft_keywords = ['leadership', 'liderazgo', 'communication', 'comunicación',
                        'problem solving', 'resolución de problemas', 'teamwork', 'trabajo en equipo',
                        'project management', 'gestión de proyectos', 'analytical', 'analítico',
                        'creative', 'creativo', 'adaptable', 'adaptabilidad']
        
        for keyword in soft_keywords:
            if keyword.lower() in text.lower():
                soft_skills.append(keyword.title())
        
        return list(set(soft_skills))
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extrae idiomas"""
        languages = []
        
        # Buscar sección de idiomas
        pattern = r'(?:LANGUAGES|IDIOMAS)(.*?)(?:CERTIFICATIONS|SKILLS|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            lang_text = match.group(1)
            candidates = re.split(r'[,\n•\-]', lang_text)
            for lang in candidates:
                lang = lang.strip()
                if len(lang) > 2 and len(lang) < 30:
                    languages.append(lang)
        
        # Buscar idiomas comunes en todo el texto
        common_langs = ['Spanish', 'Español', 'English', 'Inglés',
                       'German', 'Alemán']
        
        for lang in common_langs:
            if lang.lower() in text.lower():
                languages.append(lang)
        
        return list(set(languages))
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extrae certificaciones"""
        certs = []
        
        pattern = r'(?:CERTIFICATIONS|CERTIFICACIONES)(.*?)(?:$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            certs_text = match.group(1)
            candidates = re.split(r'\n', certs_text)
            for cert in candidates:
                cert = cert.strip()
                if len(cert) > 5:
                    certs.append(cert)
        
        return certs[:10]  # Máx 10
    
    # Métodos auxiliares de extracción
    def _extract_company_name(self, text: str) -> str:
        """Extrae nombre de empresa de un entry"""
        lines = text.split('\n')[:2]
        return lines[0].strip() if lines else ""
    
    def _extract_job_position(self, text: str) -> str:
        """Extrae posición de un entry"""
        lines = text.split('\n')
        for line in lines[:3]:
            if any(word in line.lower() for word in ['engineer', 'developer', 'manager', 'lead', 'specialist']):
                return line.strip()
        return lines[1].strip() if len(lines) > 1 else ""
    
    def _extract_duration(self, text: str) -> str:
        """Extrae duración (ej: 2021-2023)"""
        match = re.search(r'(20\d{2}|present|today|actual)', text, re.IGNORECASE)
        return match.group(0) if match else ""
    
    def _calculate_duration_years(self, duration: str) -> float:
        """Calcula años de duración"""
        years = re.findall(r'20\d{2}', duration)
        if len(years) == 2:
            try:
                return float(years[1]) - float(years[0])
            except:
                return 0.0
        return 0.0
    
    def _extract_technologies_from_entry(self, text: str) -> List[str]:
        """Extrae tecnologías de una entrada"""
        techs = []
        keywords = ['Python', 'JavaScript', 'React', 'Django', 'FastAPI', 'Docker', 'AWS', 'PostgreSQL']
        for tech in keywords:
            if tech.lower() in text.lower():
                techs.append(tech)
        return techs
    
    def _extract_achievements(self, text: str) -> List[str]:
        """Extrae logros/bullets"""
        bullets = re.findall(r'[•\-\*]\s*(.*?)(?=\n|$)', text)
        return bullets[:5] if bullets else []
    
    def _extract_institution(self, text: str) -> str:
        """Extrae institución"""
        lines = text.split('\n')
        return lines[0].strip() if lines else ""
    
    def _extract_degree(self, text: str) -> str:
        """Extrae grado/título"""
        degrees = ['Bachelor', 'Master', 'PhD', 'Diploma', 'Certificate',
                   'Licenciatura', 'Maestría', 'Doctorado', 'Diploma']
        for degree in degrees:
            if degree.lower() in text.lower():
                return degree
        return ""
    
    def _extract_field(self, text: str) -> str:
        """Extrae campo de estudio"""
        lines = text.split('\n')
        return lines[1].strip() if len(lines) > 1 else ""
    
    def _extract_year(self, text: str) -> Optional[str]:
        """Extrae año de graduación"""
        match = re.search(r'(20\d{2})', text)
        return match.group(0) if match else None
    
    def _calculate_total_years(self, experience: List[Dict[str, Any]]) -> float:
        """Calcula años totales de experiencia"""
        total = 0.0
        for job in experience:
            total += job.get('duration_years', 0.0)
        return round(total, 1)
    
    def save_to_cache(self, context: CVContext, cache_path: Optional[str] = None) -> None:
        """Guarda contexto en cache JSON"""
        path = cache_path or self.cache_path
        
        # Asegurar que existe el directorio
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        
        # Preparar datos para JSON (convertir dataclasses)
        data = {
            'name': context.name,
            'title': context.title,
            'summary': context.summary,
            'email': context.email,
            'phone': context.phone,
            'experience': context.experience,
            'education': context.education,
            'technical_skills': context.technical_skills,
            'soft_skills': context.soft_skills,
            'languages': context.languages,
            'certifications': context.certifications,
            'total_years_experience': context.total_years_experience
        }
        
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        if self.debug:
            print(f"💾 CV cacheado en: {path}")
    
    def load_from_cache(self, cache_path: Optional[str] = None) -> CVContext:
        """Carga contexto desde cache JSON con nuevo formato"""
        path = cache_path or self.cache_path
        
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Mapear desde nuevo formato JSON (con "candidate" nested)
        candidate = data.get('candidate', {})
        
        # Calcular años de experiencia desde experience
        total_years = 0.0
        experience = data.get('professional_experience', [])
        for exp in experience:
            total_years += exp.get('duration_months', 0) / 12.0
        
        # Reconstruir contexto
        context = CVContext(
            name=candidate.get('name', ''),
            title=candidate.get('current_title', ''),
            summary=candidate.get('professional_summary', ''),
            email=candidate.get('email'),
            phone=candidate.get('phone'),
            experience=experience,
            education=data.get('education', []),
            technical_skills=data.get('technical_skills', {}),
            soft_skills=data.get('soft_skills', []),
            languages=data.get('languages', []),
            certifications=data.get('certifications', []),
            total_years_experience=total_years,
            raw_text=""  # No se carga para ahorrar espacio
        )
        
        return context
    
    def get_context_as_string(self, context: Optional[CVContext] = None) -> str:
        """
        Obtiene contexto del CV como string para usar en prompts IA
        
        Args:
            context: CVContext (si no especifica, usa self.context)
        
        Returns:
            String formateado para usar en prompts
        """
        ctx = context or self.context
        if not ctx:
            raise ValueError("No CV context loaded")
        
        lines = []
        lines.append(f"CANDIDATO: {ctx.name}")
        lines.append(f"TÍTULO: {ctx.title}")
        lines.append(f"AÑOS EXPERIENCIA: {ctx.total_years_experience}")
        
        if ctx.summary:
            lines.append(f"\nRESUMEN:\n{ctx.summary}")
        
        # Información sobre autorización de trabajo (importante para preguntas de elegibilidad)
        lines.append(f"\nAUTORIZACIÓN DE TRABAJO:")
        lines.append(f"  • Residencia: Santiago, Chile")
        lines.append(f"  • Estado Legal: Legalmente autorizado para trabajar")
        lines.append(f"  • Patrocinio de Visa: NO requiere (residencia permanente)")
        
        # Información sobre CV en inglés (importante para preguntas de "résumé en inglés")
        lines.append(f"\nCURRICULUM VITAE DISPONIBLE:")
        lines.append(f"  • Versión en Español: Disponible")
        lines.append(f"  • Versión en Inglés: Disponible (puede ser proporcionada si se solicita)")
        
        # Manejar technical_skills (es dict por categorías)
        if ctx.technical_skills:
            if isinstance(ctx.technical_skills, dict):
                lines.append(f"\nSKILLS TÉCNICOS CON AÑOS:")
                for category, skills in ctx.technical_skills.items():
                    if isinstance(skills, list) and skills:
                        skill_strs = []
                        for s in skills[:8]:  # Mostrar más skills incluyendo años
                            if isinstance(s, dict):
                                skill_name = s.get('skill') or s.get('name', '')
                                years = s.get('years')
                                years_str = f" ({years} años)" if years and years > 0 else ""
                                if skill_name:
                                    skill_strs.append(skill_name + years_str)
                            else:
                                skill_strs.append(str(s))
                        skill_strs = [n for n in skill_strs if n]
                        if skill_strs:
                            lines.append(f"  • {category.replace('_', ' ').title()}: {', '.join(skill_strs)}")
        
        # Manejar soft_skills (lista de strings)
        if ctx.soft_skills:
            if isinstance(ctx.soft_skills, list):
                soft_strs = [str(s) for s in ctx.soft_skills if s]
                if soft_strs:
                    lines.append(f"\nSOFT SKILLS: {', '.join(soft_strs)}")
        
        # Manejar languages (lista de dicts con language y proficiency)
        if ctx.languages:
            if isinstance(ctx.languages, list) and ctx.languages:
                lang_strs = []
                for lang in ctx.languages:
                    if isinstance(lang, dict):
                        lang_name = lang.get('language', '')
                        proficiency = lang.get('proficiency', '')
                        if proficiency:
                            lang_strs.append(f"{lang_name} - {proficiency}")
                        else:
                            lang_strs.append(lang_name)
                    else:
                        lang_strs.append(str(lang))
                if lang_strs:
                    lines.append(f"\nIDOMAS: {', '.join(lang_strs)}")
        
        # Experience
        if ctx.experience:
            lines.append(f"\nEXPERIENCIA LABORAL:")
            for exp in ctx.experience[:5]:  # Últimos 5
                duration = exp.get('duration_months', 0) / 12.0
                lines.append(f"  • {exp['position']} en {exp['company']} ({duration:.1f} años)")
                if exp.get('technologies'):
                    lines.append(f"    Tecnologías: {','.join(exp['technologies'])}")
        
        # Education
        if ctx.education:
            lines.append(f"\nEDUCACIÓN:")
            for edu in ctx.education:
                field_str = f" - {edu['field']}" if edu.get('field') else ""
                lines.append(f"  • {edu['degree']}{field_str} en {edu['institution']}")
        
        return "\n".join(lines)


# Ejemplo de uso
if __name__ == "__main__":
    os.environ['DEBUG'] = 'True'
    
    try:
        processor = CVProcessor()
        
        # Cargar o crear contexto
        context = processor.load_or_create()
        
        print("\n=== CV CONTEXT CARGADO ===")
        print(processor.get_context_as_string(context))
    
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
